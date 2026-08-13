#!/usr/bin/env python3
"""Generate ASET 2026 submission Word document with exact format requirements.

Format:
- A4, standard margins (2.54cm)
- Title: 18pt bold centered
- Author info: 12pt left-aligned
- Body: 12pt, single spacing, left-aligned
  - Chinese: 標楷體 (DFKai-SB)
  - Latin: Times New Roman
- Section headings: 14pt bold left-aligned
- Page 1: Chinese abstract
- Pages 2-4: Research short paper
- APA 7th edition references
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent

CHINESE_FONT = "標楷體"
LATIN_FONT = "Times New Roman"

FIGURE1 = ROOT / "figure1_web_complexity_distribution.png"
FIGURE2 = ROOT / "figure2_prior_advantage_convergence.png"


def set_cell_font(run, font_cn=CHINESE_FONT, font_latin=LATIN_FONT, size=12, bold=False):
    """Set font for a run with proper CJK + Latin font handling."""
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = font_latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), font_latin)
    rfonts.set(qn("w:hAnsi"), font_latin)
    rfonts.set(qn("w:eastAsia"), font_cn)


def add_paragraph_with_font(doc, text, font_cn=CHINESE_FONT, font_latin=LATIN_FONT,
                            size=12, bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                            space_after=Pt(0), space_before=Pt(0)):
    """Add a paragraph with proper font settings."""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    set_cell_font(run, font_cn, font_latin, size, bold)
    return p


def add_heading_para(doc, text, size=14, bold=True):
    """Add a section heading (14pt bold left-aligned)."""
    return add_paragraph_with_font(doc, text, size=size, bold=bold,
                                    alignment=WD_ALIGN_PARAGRAPH.LEFT,
                                    space_before=Pt(6), space_after=Pt(3))


def add_body_para(doc, text, size=12):
    """Add a body paragraph (12pt, left-aligned, single spacing)."""
    return add_paragraph_with_font(doc, text, size=size, bold=False,
                                    alignment=WD_ALIGN_PARAGRAPH.LEFT)


def add_page_break(doc):
    """Add a page break."""
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx.enum.text.WD_BREAK.PAGE) if hasattr(docx, "enum") else None
    from docx.enum.text import WD_BREAK
    run.add_break(WD_BREAK.PAGE)


def build_document():
    doc = Document()

    # Set A4 page size and standard margins
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # Set default style font
    style = doc.styles["Normal"]
    style.font.name = LATIN_FONT
    style.font.size = Pt(12)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), CHINESE_FONT)
    rfonts.set(qn("w:ascii"), LATIN_FONT)
    rfonts.set(qn("w:hAnsi"), LATIN_FONT)

    # ========== PAGE 1: CHINESE ABSTRACT ==========

    # Title (18pt bold centered)
    add_paragraph_with_font(doc,
        "AI 輔助地球科學數位作品的縱向發展：先備數位經驗、早期優勢與組間趨同",
        size=18, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=Pt(12))

    # Author info (12pt left-aligned)
    add_paragraph_with_font(doc, "陳達毅¹²", size=12, bold=False,
                            alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph_with_font(doc, "¹中央氣象署地震測報中心", size=12,
                            alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph_with_font(doc, "²臺北市立大學地球環境暨生物資源學系", size=12,
                            alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph_with_font(doc, "oceanicdayi@gmail.com", size=12,
                            alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(12))

    # Abstract heading (14pt bold centered)
    add_paragraph_with_font(doc, "摘要", size=14, bold=True,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))

    # Abstract body (12pt left-aligned, single spacing)
    abstract_zh = (
        "本研究探討學生在連續地球科學課程中使用生成式 AI 與網頁部署工具後，"
        "數位作品如何發展，並分析先備經驗與後續表現的關聯。"
        "研究以臺北市立大學地球物理與地震學課程學生網頁為分析單位，"
        "以程式量化媒體豐富度、排版架構、可見文本字元數、"
        "專有名詞使用及啟發式批判反思訊號。"
        "指導老師對 53 份作品兩次盲評，"
        "媒體與排版尺規 κ 為 0.75 與 0.66，達 substantial agreement。"
        "程式分析跨學期配對（n=8）顯示媒體由 1.75 升至 3.50，"
        "文本由 1,983 增至 3,693 字元（+86%），"
        "反思訊號由 6.54% 升至 10.73%。"
        "人工評分（n=11）亦顯示媒體由 2.45 升至 3.45、"
        "排版由 2.64 升至 3.09。"
        "第一次作業先備組文本比新加入組多 1,743 字、"
        "專有名詞多 29 次（δ=0.36–0.43）但未達顯著，"
        "人工評分先備組媒體亦較高（2.21 vs 1.80）。"
        "至期末，程式分析組間差距縮小（δ 0.354→0.092），"
        "但人工評分先備組媒體仍較高（3.38 vs 2.20）。"
        "結果顯示先備經驗與早期內容優勢相關，"
        "共同 AI 輔助製作伴隨部分指標趨近，"
        "惟自動化與人工在收斂趨勢上有差異。"
        "因樣本小、非隨機且作業不同，屬探索性。"
    )
    add_body_para(doc, abstract_zh)

    # Keywords
    add_paragraph_with_font(doc, "", size=6)
    add_body_para(doc, "關鍵詞：生成式 AI、地球科學教育、數位作品、縱向研究、數位敘事")

    # ========== PAGE BREAK ==========
    add_page_break(doc)

    # ========== PAGES 2-4: RESEARCH SHORT PAPER ==========

    # Title (18pt bold centered)
    add_paragraph_with_font(doc,
        "AI 輔助地球科學數位作品的縱向發展：先備數位經驗、早期優勢與組間趨同",
        size=18, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(8))

    # Author info
    add_paragraph_with_font(doc, "陳達毅¹²", size=12, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph_with_font(doc, "¹中央氣象署地震測報中心", size=12,
                            alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph_with_font(doc, "²臺北市立大學地球環境暨生物資源學系", size=12,
                            alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph_with_font(doc, "oceanicdayi@gmail.com", size=12,
                            alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(6))

    # --- 研究目的 ---
    add_heading_para(doc, "研究目的")

    add_body_para(doc,
        "生成式 AI 已快速進入大學教學，但其教育效果不應只以「是否使用 AI」"
        "或單次測驗成績判斷（Kasneci et al., 2023; UNESCO, 2023）。當學生利用"
        " ChatGPT、Gemini 等工具協助整理科學內容，再以 GitHub Pages、Hugging Face"
        " 或其他網頁平台將學習成果轉化為可瀏覽、可互動的數位作品時，學生同時進行"
        "學科內容重組、數位敘事與資訊工具整合（Robin, 2016; Tyagi, 2023）。"
        "這類數位成品保留了學習過程留下的結構性痕跡，因此可作為觀察學生如何使用"
        " AI 進行知識表達的另一種證據來源。")

    add_body_para(doc,
        "本研究以連續兩學期的「地球物理」與「地震學」課程為場域，學生在課程中"
        "持續使用生成式 AI 與網頁部署工具製作學習成果。研究不直接把網頁精美程度"
        "視為學科成績，而是將作品視為 digital artifacts（Penuel et al., 2022），"
        "分析其呈現形式、內容量、學科語言及反思訊號的變化。本研究提出三個研究問題：")

    add_body_para(doc, "1. RQ1：同一批學生的數位作品，在連續兩學期中如何改變？")
    add_body_para(doc, "2. RQ2：已有前一學期數位作品紀錄的學生，是否在下一學期第一次網頁作業中呈現較強的內容建構表現？")
    add_body_para(doc, "3. RQ3：若存在早期差異，該差異是否在一學期共同的 AI 輔助網頁製作經驗後縮小？")

    # --- 研究重要性 ---
    add_heading_para(doc, "研究重要性")

    add_body_para(doc,
        "此研究的重要性在於：相較於只比較期末成績，本研究嘗試建立一套可重複執行的"
        " digital-artifact analysis 方法，將學生作品本身轉化為可分析的教育研究資料"
        "（Krajcik & Blumenfeld, 2006; Penuel et al., 2022），同時保留對小樣本、"
        "非隨機設計與自動評量限制的謹慎解讀。")

    # --- 研究方法 ---
    add_heading_para(doc, "研究方法")

    add_body_para(doc,
        "研究資料來自兩個連續的大學地球科學課程，公開資料均以一次性匿名代碼呈現。"
        "分析分成三個層次：（1）跨學期縱向配對 13 人，其中 8 人符合高信度條件，用於 RQ1；"
        "（2）下學期第一次網頁作業組間比較：先備組 14 人 vs 新加入組 6 人，用於 RQ2；"
        "（3）第一次作業至期末變化 18 人（先備組 13、新加入組 5），用於 RQ3。"
        "「有前一學期作品紀錄」僅作為先備數位作品經驗的代理變項。")

    add_body_para(doc,
        "研究以 Python 與 BeautifulSoup 解析 HTML，量化五項指標：媒體豐富度（1–4）、"
        "排版架構（1–4）、可見文本字元數、專有名詞使用（絕對次數與每千字密度）、"
        "及啟發式批判反思訊號（Newell et al., 2011）。上述指標描述作品特徵，"
        "不直接衡量地球科學內容正確性。媒體與排版為規則式操作化尺規"
        "（Jonassen, 2006）；批判反思訊號亦非經驗證的心理量表。")

    add_body_para(doc,
        "為檢驗尺規的評分穩定性，指導老師對 57 份作品進行兩次盲評"
        "（間隔約 30 分鐘，順序經亂序處理）。4 份因登入牆排除，最終納入 53 份。"
        "此設計檢驗 test-retest reliability（Cohen, 1960; Landis & Koch, 1977），"
        "而非評分者間一致性。兩次評分的 quadratic weighted Cohen's κ 分別為"
        " 0.75（媒體）與 0.66（排版），均達 substantial agreement。"
        "人工共識與自動化評分的比較顯示，媒體 exact agreement=0.264、"
        "排版=0.509，反映自動規則式判定與人工判斷間存在系統性差異。")

    add_body_para(doc,
        "統計分析使用描述統計、個別配對軌跡、Cliff's delta（Cliff, 1993）、"
        "精確秩排列檢定與 Fisher exact test。由於樣本小且非隨機分組，"
        "統計檢定主要用於描述分布與效果方向，而非建立因果結論。")

    # --- 研究結果與討論 ---
    add_heading_para(doc, "研究結果與討論")

    add_body_para(doc,
        "1. 連續課程中的數位作品發展：在 8 位高信度配對學生中，媒體豐富度由"
        " 1.75 提升至 3.50（6 人提升、無人下降），可見文本由 1,983 增至 3,693 字元"
        "（7 人增加），批判反思訊號由 6.54% 增至 10.73%（6 人增加）。"
        "排版架構大致持平（3.00→2.88），可能存在尺規天花板。"
        "專有名詞密度 4 人增加、4 人下降，顯示篇幅增加不必然伴隨更高的術語集中度。")

    # Insert Figure 1
    if FIGURE1.exists():
        p_fig1 = doc.add_paragraph()
        p_fig1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_fig1 = p_fig1.add_run()
        run_fig1.add_picture(str(FIGURE1), width=Cm(16))
        add_paragraph_with_font(doc,
            "圖一 8 位高信度配對學生上、下學期媒體豐富度與排版架構評分分布",
            size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))

    add_body_para(doc,
        "2. 先備數位作品經驗的早期優勢：先備組（n=14）的平均可見文本比新加入組"
        "（n=6）多約 1,743 字元，專有名詞絕對次數多 29 次，主題涵蓋多 1.36/11。"
        "三項內容指標方向一致（Cliff's δ=0.36–0.43），但精確檢定未達 p < .05，"
        "屬探索性證據。媒體與排版中位數兩組相同（均為 3.00），"
        "顯示先備優勢在內容建構而非結構面。")

    add_body_para(doc,
        "3. 學期中的組間趨同：先備組媒體豐富度由 3.23 增至 3.54（+0.31），"
        "新加入組由 2.80 增至 3.40（+0.60）。批判反思訊號方面，先備組由 9.72% 增至"
        " 12.67%（+2.95pp），新加入組由 7.38% 增至 13.98%（+6.60pp）。"
        "至期末，組間效果量大幅縮小：媒體 δ 由 0.354 降至 0.092，"
        "反思訊號 δ 由 0.262 降至 0.015，新加入組甚至超過先備組。"
        "排版兩組皆無變化（尺規天花板）。可見文本先備組由 5,030 降至 3,802，"
        "新加入組由 3,179 升至 4,227，方向反轉，反映期末作品改採互動呈現而非純文字。"
        "此與「先備經驗的早期差距在共同經驗累積後逐步縮小」之描述一致，"
        "但因兩次作業目的不同，不能視為等值前後測。")

    # Insert Figure 2
    if FIGURE2.exists():
        p_fig2 = doc.add_paragraph()
        p_fig2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_fig2 = p_fig2.add_run()
        run_fig2.add_picture(str(FIGURE2), width=Cm(16))
        add_paragraph_with_font(doc,
            "圖二 先備組與新加入組由第一次網頁作業至期末的媒體豐富度與排版架構平均軌跡",
            size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))

    add_body_para(doc,
        "4. 人工評分結果：程式分析與人工評分的交叉比較。"
        "指導老師兩次盲評共識分數提供與程式自動評分交叉比對的機會。"
        "11 位上下學期均有期末作品的學生，人工評分媒體由 2.45 升至 3.45（+1.00），"
        "排版由 2.64 升至 3.09（+0.45），與程式分析「媒體明顯提升」方向一致，"
        "但人工評分亦觀察到排版提升，而程式分析因尺規天花板呈現持平。"
        "第一次作業中，人工評分先備組媒體（2.21）高於新加入組（1.80），"
        "與程式分析方向一致。至期末，人工評分先備組媒體（3.38）"
        "仍明顯高於新加入組（2.20），差距 +1.18，"
        "與程式分析顯示的組間收斂趨勢不一致。"
        "此差異可能源自程式以 HTML 標籤判定，"
        "而人工以頁面實際呈現為準：部分新加入組學生使用 React/SPA 框架，"
        "程式因靜態 HTML 稀疏而給高分，但人工評分認為實際互動深度有限。"
        "此交叉比對顯示，自動化與人工評分在「個別成長方向」上大致一致，"
        "但在「組間收斂趨勢」上存在分歧，"
        "呼應前述自動化與人工一致率偏低（媒體 exact=0.26）的發現。")

    add_body_para(doc,
        "5. AI 融入 STEM 之教學實踐：課程以核心地科理論與數學推演為學理基礎，"
        "並強化科技與工程實作。學生架設震測儀器、組裝 Raspberry Pi 物聯網地震儀，"
        "並使用 ChatGPT 與 Gemini 作為程式除錯助教。兩學期皆以「數位學習歷程網頁」"
        "取代期末考，運用 GitHub Pages 與 Hugging Face 部署。學生將學習成果轉化為"
        "公開數位作品，其產出的媒體豐富度與可見字數均有成長，期末作品中的"
        "批判思考與反思訊號比例亦有所提升。")

    add_body_para(doc,
        "6. 討論與限制：本研究顯示，學生數位作品最一致的變化出現在媒體整合與內容產出，"
        "而「如何使用 AI 與網頁工具表達科學內容」本身可能是一種可累積的數位經驗"
        "（Mishra & Koehler, 2006）。程式分析顯示新加入組增幅較大，組間差距至期末縮小；"
        "然而人工評分結果呈現不同圖像：先備組期末媒體（3.38）仍明顯高於新加入組（2.20），"
        "差距並未收斂。此分歧可能反映自動化尺規對 React/SPA 框架的過度評價，"
        "亦提示「組間趨同」此一結論需以更嚴謹的評量方式驗證。"
        "提示 AI 輔助網頁製作可能為先備經驗較弱學生提供追趕空間（Wang et al., 2024），"
        "但此結果不應被解讀為 AI 已證實能消除學習落差。")

    add_body_para(doc,
        "本研究有四項主要限制：第一，樣本小且非隨機分組。第二，「前一學期作品紀錄」"
        "只是先備經驗代理變項。第三，不同時間點作業目的不同，不能視為標準化前後測。"
        "第四，數位作品尺規與批判反思訊號主要由規則式方法操作化，指導老師兩次評分"
        "達 substantial agreement（κ=0.75/0.66），但自動化與人工一致率偏低"
        "（媒體 exact=0.26），且批判反思訊號尚未經人工內容分析驗證，"
        "應視為 heuristic critical-reflection signal。後續研究應納入獨立評分者"
        "以建立評分者間信度。整體而言，本研究支持將學生數位作品作為觀察 AI 融入"
        "科學教育歷程的資料來源（Penuel et al., 2022），為後續更大樣本的縱向研究"
        "建立分析基礎。")

    # --- 參考文獻 ---
    add_heading_para(doc, "參考文獻")

    refs = [
        "Cliff, N. (1993). Dominance statistics: Ordinal analyses to answer ordinal questions. Psychological Bulletin, 114(3), 494–509.",
        "Cohen, J. (1960). A coefficient of agreement for nominal scales. Educational and Psychological Measurement, 20(1), 37–46.",
        "Jonassen, D. H. (2006). On the role of concepts in understanding and identifying learning artifacts. Educational Technology Research and Development, 54(2), 177–189.",
        "Kasneci, E., Sessler, K., Küchemann, S., Bannert, M., Dementieva, D., Fischer, F., Gasser, U., Groh, G., Günnemann, S., Hüllermeier, E., Krusche, S., Kutyniok, G., Lermann, T., Mittelstädt, A., Plonner, M., Ratklies, L., Scheuermann, M., & Kasneci, G. (2023). ChatGPT for good? On opportunities and challenges of large language models for education. Learning and Individual Differences, 103, 102274.",
        "Krajcik, J. S., & Blumenfeld, P. C. (2006). Project-based learning. In R. K. Sawyer (Ed.), The Cambridge Handbook of the Learning Sciences (pp. 317–334). Cambridge University Press.",
        "Landis, J. R., & Koch, G. G. (1977). The measurement of observer agreement for categorical data. Biometrics, 33(1), 159–174.",
        "Mishra, P., & Koehler, M. J. (2006). Technological pedagogical content knowledge: A framework for teacher knowledge. Teachers College Record, 108(6), 1017–1054.",
        "Newell, G. E., Beach, R., Smith, J., & VanDerHeide, J. (2011). Teaching and learning argumentative reading and writing: A review of research. Reading Research Quarterly, 46(3), 273–304.",
        "Penuel, W. R., Coburn, C. E., & Gallagher, D. J. (2022). Out of the margins: The promise of practice-based approaches to research on teaching and learning. Educational Researcher, 51(2), 131–141.",
        "Robin, B. R. (2016). Digital storytelling: A powerful technology tool for the 21st century classroom. Theory Into Practice, 47(3), 220–228.",
        "Tyagi, S. (2023). Adoption of generative AI in higher education: A narrative review. Journal of Educational Technology Systems, 52(3), 302–318.",
        "UNESCO. (2023). Guidance for generative AI in education and research. UNESCO Publishing.",
        "Wang, S., Lin, Y., & Chiu, C. (2024). Exploring the impact of generative AI on student learning outcomes in STEM education. Computers & Education, 211, 104955.",
    ]
    for ref in refs:
        p_ref = doc.add_paragraph()
        p_ref.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_ref.paragraph_format.space_after = Pt(3)
        p_ref.paragraph_format.line_spacing = 1.0
        # Hanging indent for references
        p_ref.paragraph_format.left_indent = Cm(1.27)
        p_ref.paragraph_format.first_line_indent = Cm(-1.27)
        run_ref = p_ref.add_run(ref)
        set_cell_font(run_ref, size=12)

    # Save
    out_docx = ROOT / "ASET_2026_審查用短文_v3.docx"
    doc.save(str(out_docx))
    print(f"[saved] {out_docx.name}")

    # Convert to PDF
    import shutil
    try:
        from docx2pdf import convert
        convert(str(out_docx), str(ROOT / "ASET_2026_審查用短文_v3.pdf"))
        print(f"[saved] ASET_2026_審查用短文.pdf")
    except Exception as e:
        print(f"[warn] PDF conversion failed: {e}")


if __name__ == "__main__":
    import docx
    build_document()